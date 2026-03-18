from __future__ import annotations

import csv
import io
import json
import os
import re
import shutil
import tempfile
import traceback
import zipfile
from collections import Counter, defaultdict
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

from docx import Document
from fastapi import Body, FastAPI, File, UploadFile
from fastapi.responses import HTMLResponse, JSONResponse, Response
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from openai import OpenAI
from openpyxl import load_workbook
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer
from starlette.requests import Request

# Optional OCR
try:
    from PIL import Image  # type: ignore
except Exception:
    Image = None

try:
    import pytesseract  # type: ignore
except Exception:
    pytesseract = None

try:
    from pdf2image import convert_from_path  # type: ignore
except Exception:
    convert_from_path = None


app = FastAPI(title="Tender Pack Summary")

app.mount("/static", StaticFiles(directory="app/static"), name="static")
templates = Jinja2Templates(directory="templates")

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "").strip()
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-4o-mini").strip() or "gpt-4o-mini"
OPENAI_TIMEOUT_SEC = float(os.getenv("OPENAI_TIMEOUT_SEC", "20"))
openai_client = OpenAI(api_key=OPENAI_API_KEY) if OPENAI_API_KEY else None

MAX_FILE_BYTES = 300 * 1024 * 1024
MAX_FILES_TO_PROCESS = 20000
MAX_ZIP_MEMBERS = 30000
MAX_ZIP_RECURSION = 3
EXTRACT_WORKERS = 8

PDF_MAX_PAGES = 30
PDF_OCR_MAX_PAGES = 8
DOCX_MAX_PARAS = 1200
DOCX_MAX_TABLES = 80
DOCX_MAX_TABLE_ROWS = 500
XLSX_MAX_SHEETS = 24
XLSX_MAX_ROWS = 80
CSV_MAX_ROWS = 80

IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tif", ".tiff"}
DRAWING_EXTS = {".dwg", ".dxf"}
SPREADSHEET_EXTS = {".xlsx", ".xls", ".csv"}
DOC_EXTS = {".docx", ".doc"}
PDF_EXTS = {".pdf"}

IGNORED_FILENAMES = {".ds_store", "thumbs.db", "desktop.ini"}
IGNORED_SUFFIXES = {".tmp", ".bak", ".lock"}

KEYWORDS = {
    "boq": ["boq", "bill", "bill of quantities", "pricing schedule", "schedule of rates", "sor", "rates", "pricing"],
    "register": ["register", "drawing register", "document register", "issue register", "transmittal"],
    "addenda": ["addendum", "addenda", "clarification", "query response", "rfi response", "tender query", "tq"],
    "prelims": ["prelim", "prelims", "preliminary"],
    "specs": ["spec", "specification", "employer requirements", "employer's requirements", "works information", "scope"],
    "forms": ["form", "tender form", "declaration", "questionnaire", "appendix", "submission", "itt", "invitation to tender"],
}

ESTIMATOR_KEYWORDS = [
    "asbestos", "acm", "soft strip", "strip out", "demolition", "temporary works", "propping",
    "party wall", "working hours", "out of hours", "noise", "dust", "vibration",
    "traffic management", "permit", "permits", "consent", "licence", "license",
    "waste", "recycling", "segregation", "muck away", "skip", "haulage", "crushing", "arisings",
    "water", "electric", "gas", "services", "live", "isolation", "disconnect", "diversion",
    "section 61", "access", "logistics", "hoarding", "scaffold", "crane", "lift",
    "phasing", "sequence", "sequencing", "retention", "bond", "warranty", "insurance",
    "liquidated damages", "ld", "lad", "penalty", "tender", "submission", "deadline",
]

RISK_BUCKETS: dict[str, list[str]] = {
    "Asbestos / hazardous materials": ["asbestos", "acm", "hazardous", "lead paint", "silica"],
    "Temporary works / propping": ["temporary works", "propping", "needling", "backpropping", "sequencing"],
    "Party wall / adjacent structures": ["party wall", "adjoining", "adjacent", "third party", "neighbour", "neighbor"],
    "Traffic management / access": ["traffic management", "road closure", "lane closure", "delivery", "access", "logistics", "banksman"],
    "Noise / dust / vibration": ["noise", "dust", "vibration", "monitoring", "suppression", "section 61"],
    "Services / isolations": ["isolation", "isolations", "live", "disconnect", "disconnection", "divert", "diversion", "electric", "gas", "water", "drainage"],
    "Waste / crushing / segregation": ["waste", "segregation", "recycling", "crushing", "arisings", "haulage", "skip", "muck away"],
    "Permits / licences": ["permit", "licence", "license", "consent", "section 61"],
    "Working hours / constraints": ["working hours", "site hours", "out of hours", "weekend", "night works"],
    "Liquidated damages / penalties": ["liquidated damages", "ld", "lds", "lad", "lads", "penalty"],
}

DATE_PATTERNS = [
    r"\b(\d{1,2}\s+(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\s+\d{2,4})\b",
    r"\b(\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4})\b",
]

REQ_STRICT_RE = re.compile(r"\b(must|shall|required|mandatory|as a minimum|minimum of|no later than)\b", re.I)
REQ_LOOSE_RE = re.compile(r"\b(should|please|requested|provide|submit|include|confirm)\b", re.I)
COMMERCIAL_SIGNAL_RE = re.compile(r"(£\s?\d|%\b|\bweeks?\b|\bmonths?\b|\b\d{1,2}[:.]\d{2}\b)", re.I)
SENT_SPLIT_RE = re.compile(r"(?<=[\.\!\?])\s+|\n+")
IRRELEVANT_DOC_HINTS = [
    "breeam", "credit", "wat 01", "assessor", "calculator", "guidance",
    "performance levels", "this document represents guidance",
]
BAD_FRAGMENT_RE = re.compile(
    r"(\b\d+(\.\d+){1,}\b|"
    r"\b(employer'?s agent|project manager|quantity surveyor)\b|"
    r"\bappendix\b|\bclause\b|\bschedule\b\s+\d|\bsection\b\s+\d)",
    re.I,
)

EMAIL_RE = re.compile(r"\b([A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,})\b", re.I)


# ---------------- generic helpers ----------------
def _normalize_line(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "")).strip()


def _clean_text(s: str) -> str:
    s = (s or "").replace("\x00", "")
    s = re.sub(r"(\w)-\n(\w)", r"\1\2", s)
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    s = re.sub(r"\n{3,}", "\n\n", s)
    s = re.sub(r"(?<!\n)\n(?!\n)", " ", s)
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()


def _has_any(haystack: str, words: list[str]) -> bool:
    h = (haystack or "").lower()
    return any(w in h for w in words)


def _safe_json_loads(s: str) -> dict[str, Any] | None:
    try:
        data = json.loads(s)
        return data if isinstance(data, dict) else None
    except Exception:
        return None


def _trim_list(items: Any, limit: int = 10, width: int = 320) -> list[str]:
    if not isinstance(items, list):
        return []
    out: list[str] = []
    seen: set[str] = set()
    for x in items:
        if not isinstance(x, str):
            continue
        x2 = _normalize_line(x)[:width]
        if not x2:
            continue
        k = x2.lower()
        if k in seen:
            continue
        seen.add(k)
        out.append(x2)
        if len(out) >= limit:
            break
    return out


def _looks_irrelevant(text: str) -> bool:
    tl = (text or "").lower()
    return any(h in tl for h in IRRELEVANT_DOC_HINTS)


def _looks_like_schedule_row(s: str) -> bool:
    s2 = _normalize_line(s)
    if not s2:
        return True
    if s2.count("|") >= 2:
        return True
    tokens = re.findall(r"\b[A-Z]{1,4}[-_ ]?\d{1,6}[A-Z]?\b", s2)
    if len(tokens) >= 5:
        return True
    digits = sum(1 for c in s2 if c.isdigit())
    if digits >= 18 and len(s2) < 170:
        return True
    return False


def _is_gibberish_line(s: str) -> bool:
    s2 = _normalize_line(s)
    if len(s2) < 12:
        return True
    if _looks_like_schedule_row(s2):
        return True
    letters = [c for c in s2 if c.isalpha()]
    if letters:
        upper_ratio = sum(1 for c in letters if c.isupper()) / max(1, len(letters))
        if upper_ratio > 0.78 and len(s2) > 50:
            return True
    return False


def _looks_bad_fragment(s: str) -> bool:
    s2 = _normalize_line(s)
    if not s2:
        return True
    if BAD_FRAGMENT_RE.search(s2) and len(s2) < 240:
        return True
    return False


def _clean_candidate_sentence(s: str, width: int = 360) -> str | None:
    s2 = _normalize_line(s)
    if not s2:
        return None
    if _looks_irrelevant(s2) or _looks_like_schedule_row(s2) or _is_gibberish_line(s2) or _looks_bad_fragment(s2):
        return None
    if s2[:1].islower():
        return None
    if len(s2.split()) < 7:
        return None
    return s2[:width]


def _split_sentences(text: str) -> list[str]:
    if not text:
        return []
    return [p.strip() for p in SENT_SPLIT_RE.split(text) if p and p.strip()]


def _sentences_around(text: str, idx: int, max_sentences: int = 2) -> str:
    if not text:
        return ""
    parts = SENT_SPLIT_RE.split(text)
    offsets: list[tuple[int, int, str]] = []
    pos = 0
    for p in parts:
        p2 = p.strip()
        if not p2:
            pos += len(p) + 1
            continue
        start = text.find(p2, pos)
        if start == -1:
            start = pos
        end = start + len(p2)
        offsets.append((start, end, p2))
        pos = end

    s_i = 0
    for i, (a, b, _) in enumerate(offsets):
        if a <= idx <= b:
            s_i = i
            break

    chosen: list[str] = []
    for j in range(max(0, s_i), min(len(offsets), s_i + max_sentences)):
        cleaned = _clean_candidate_sentence(offsets[j][2], width=360)
        if cleaned:
            chosen.append(cleaned)
    return " ".join(chosen).strip()


def _sentence_score(s: str) -> int:
    s_l = s.lower()
    score = 0

    if COMMERCIAL_SIGNAL_RE.search(s):
        score += 6

    for kw in ESTIMATOR_KEYWORDS:
        if kw in s_l:
            score += 3

    if REQ_STRICT_RE.search(s):
        score += 2
    if REQ_LOOSE_RE.search(s):
        score += 1

    if BAD_FRAGMENT_RE.search(s):
        score -= 6

    if len(s) < 90:
        score -= 2
    if s.count(" ") < 12:
        score -= 3

    return score


def _dedup_keep_best(items: list[str], limit: int, width: int = 320) -> list[str]:
    scored: list[tuple[int, str]] = []
    seen: set[str] = set()
    for x in items:
        x2 = _clean_candidate_sentence(x, width=width)
        if not x2:
            continue
        k = x2.lower()
        if k in seen:
            continue
        seen.add(k)
        scored.append((_sentence_score(x2), x2))
    scored.sort(key=lambda x: x[0], reverse=True)
    return [x[1] for x in scored[:limit]]


def _is_supported_upload(filename: str) -> bool:
    base = Path(filename).name.lower()
    if base in IGNORED_FILENAMES:
        return False
    if base.startswith("~$"):
        return False
    if Path(filename).suffix.lower() in IGNORED_SUFFIXES:
        return False
    return True


def _valid_date_candidate(s: str) -> bool:
    s2 = (s or "").strip()
    if not s2:
        return False

    match_year = re.search(r"\b(20\d{2})\b", s2)
    if match_year:
        y = int(match_year.group(1))
        if y < 2024 or y > 2035:
            return False

    if re.fullmatch(r"\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4}", s2):
        parts = re.split(r"[\/\-]", s2)
        if len(parts) == 3:
            year = parts[2]
            if len(year) == 2:
                y = int(year)
                if y < 24:
                    return False
            else:
                y = int(year)
                if y < 2024 or y > 2035:
                    return False

    return True


def _extract_date_candidates_from_text(text: str) -> list[str]:
    dates: set[str] = set()
    for pat in DATE_PATTERNS:
        for m in re.finditer(pat, text or "", flags=re.IGNORECASE):
            if _valid_date_candidate(m.group(1)):
                dates.add(m.group(1))
    return sorted(dates)[:30]


# ---------------- classification ----------------
def classify_file(p: Path, display: str | None = None) -> str:
    ext = p.suffix.lower()
    name = p.name.lower()
    full = (display or p.as_posix()).lower()

    if ext in DRAWING_EXTS:
        return "drawings"

    if ext in IMAGE_EXTS:
        return "photos"

    if ext in SPREADSHEET_EXTS:
        if _has_any(full, KEYWORDS["register"]) or _has_any(name, KEYWORDS["register"]):
            return "registers"
        if _has_any(full, KEYWORDS["boq"]) or _has_any(name, KEYWORDS["boq"]):
            return "boq"
        return "spreadsheets"

    if ext in DOC_EXTS:
        if _has_any(full, KEYWORDS["forms"]) or _has_any(name, KEYWORDS["forms"]):
            return "forms"
        if _has_any(full, KEYWORDS["prelims"]) or _has_any(name, KEYWORDS["prelims"]):
            return "prelims"
        if _has_any(full, KEYWORDS["specs"]) or _has_any(name, KEYWORDS["specs"]):
            return "specs"
        return "documents"

    if ext in PDF_EXTS:
        if _has_any(full, KEYWORDS["addenda"]) or _has_any(name, KEYWORDS["addenda"]):
            return "addenda"
        if _has_any(full, KEYWORDS["boq"]) or _has_any(name, KEYWORDS["boq"]):
            return "boq"
        if _has_any(full, KEYWORDS["register"]) or _has_any(name, KEYWORDS["register"]):
            return "registers"
        if _has_any(full, KEYWORDS["prelims"]) or _has_any(name, KEYWORDS["prelims"]):
            return "prelims"
        if _has_any(full, KEYWORDS["specs"]) or _has_any(name, KEYWORDS["specs"]):
            return "specs"
        if _has_any(full, KEYWORDS["forms"]) or _has_any(name, KEYWORDS["forms"]):
            return "forms"
        if any(h in full for h in ["drawing", "drg", "dwg", "ga", "elev", "section"]) or any(h in name for h in ["drawing", "drg", "dwg", "ga", "elev", "section"]):
            return "drawings"
        return "pdfs"

    return "other"


def _file_priority_score(display: str, category: str) -> int:
    s = (display or "").lower()
    score = 0

    priority_words = {
        130: ["invitation to tender", "itt", "tender return", "instructions to tenderers"],
        120: ["prelim", "prelims"],
        115: ["employer requirements", "employer's requirements", "works information", "scope", "specification", "spec"],
        110: ["pricing schedule", "bill of quantities", "boq", "schedule of rates"],
        100: ["clarification", "addendum", "addenda", "query response", "rfi"],
        95: ["programme", "program", "gantt"],
        90: ["form of tender", "submission", "declaration", "appendix"],
        85: ["document register", "drawing register", "register"],
    }

    for pts, words in priority_words.items():
        if any(w in s for w in words):
            score += pts

    category_bonus = {
        "prelims": 60,
        "specs": 58,
        "forms": 55,
        "addenda": 52,
        "boq": 50,
        "registers": 42,
        "pdfs": 28,
        "documents": 24,
        "spreadsheets": 15,
        "drawings": 8,
        "photos": 6,
    }
    score += category_bonus.get(category, 0)

    parts = s.split("/")
    if len(parts) > 1:
        folder_text = " ".join(parts[:-1])
        if any(w in folder_text for w in ["tender", "submission", "commercial", "prelim", "spec", "pricing", "clarification"]):
            score += 15

    return score


# ---------------- zip handling ----------------
def safe_extract_zip(zip_path: Path, extract_to: Path, max_files: int = MAX_ZIP_MEMBERS) -> list[Path]:
    extracted: list[Path] = []
    base = extract_to.resolve()

    with zipfile.ZipFile(zip_path, "r") as z:
        members = z.infolist()
        if len(members) > max_files:
            raise ValueError(f"Too many files in ZIP ({len(members)}). Limit is {max_files}.")

        for m in members:
            if m.is_dir():
                continue

            zname = (m.filename or "").replace("\\", "/").lstrip("/")
            parts = [p for p in zname.split("/") if p not in ("", ".", "..")]
            safe_name = "/".join(parts) if parts else Path(m.filename).name

            if not _is_supported_upload(safe_name) and not safe_name.lower().endswith(".zip"):
                continue

            out_path = (extract_to / safe_name).resolve()
            if not str(out_path).startswith(str(base)):
                raise ValueError("Unsafe ZIP: path traversal detected.")

            out_path.parent.mkdir(parents=True, exist_ok=True)
            with z.open(m, "r") as src, open(out_path, "wb") as dst:
                shutil.copyfileobj(src, dst)

            extracted.append(out_path)

    return extracted


def _recursive_extract_zip(zip_path: Path, extract_to: Path, depth: int = 0) -> list[Path]:
    if depth > MAX_ZIP_RECURSION:
        return []

    out: list[Path] = []
    extracted = safe_extract_zip(zip_path, extract_to, max_files=MAX_ZIP_MEMBERS)

    for p in extracted:
        if p.suffix.lower() == ".zip":
            nested_dir = p.parent / f"{p.stem}_unzipped"
            nested_dir.mkdir(parents=True, exist_ok=True)
            try:
                out.extend(_recursive_extract_zip(p, nested_dir, depth + 1))
            except Exception:
                continue
        else:
            out.append(p)

    return out


# ---------------- OCR ----------------
def _ocr_image_file(path: Path) -> str:
    if not Image or not pytesseract:
        return ""
    try:
        img = Image.open(path)
        return _clean_text(pytesseract.image_to_string(img))
    except Exception:
        return ""


def _ocr_pdf_file(path: Path, max_pages: int = PDF_OCR_MAX_PAGES) -> str:
    if not convert_from_path or not pytesseract:
        return ""
    try:
        images = convert_from_path(str(path), first_page=1, last_page=max_pages, dpi=200)
        parts: list[str] = []
        for img in images[:max_pages]:
            txt = pytesseract.image_to_string(img)
            if txt and txt.strip():
                parts.append(txt)
        return _clean_text("\n".join(parts))
    except Exception:
        return ""


# ---------------- extraction ----------------
def guess_revision(filename: str) -> str | None:
    s = (filename or "").upper()
    m = re.search(r"(?:REV[\s_\-]*)?([PC]\d{2,3})\b", s)
    return m.group(1) if m else None


def guess_drawing_number(filename: str) -> str | None:
    s = Path(filename).stem.upper()
    m = re.search(r"\b([A-Z]{1,4}[-_ ]?\d{2,6})\b", s)
    if m:
        return m.group(1).replace(" ", "-").replace("_", "-")
    return None


def _extract_requirements(text: str, max_lines: int = 22) -> dict[str, list[str]]:
    if not text:
        return {"strict": [], "loose": []}

    sentences = _split_sentences(text)
    strict_raw: list[str] = []
    loose_raw: list[str] = []

    for s in sentences:
        s2 = _clean_candidate_sentence(s, width=380)
        if not s2:
            continue
        if _sentence_score(s2) <= 0:
            continue

        if REQ_STRICT_RE.search(s2):
            strict_raw.append(s2)
        elif REQ_LOOSE_RE.search(s2):
            loose_raw.append(s2)

    return {
        "strict": _dedup_keep_best(strict_raw, max_lines, 340),
        "loose": _dedup_keep_best(loose_raw, max_lines, 340),
    }


def extract_pdf_info(path: Path, max_pages: int = PDF_MAX_PAGES) -> dict[str, Any]:
    info: dict[str, Any] = {
        "pages": None,
        "keyword_hits": {},
        "date_candidates": [],
        "snippet": "",
        "text_len": 0,
        "requirements": {"strict": [], "loose": []},
        "ocr_used": False,
    }

    try:
        reader = PdfReader(str(path))
        info["pages"] = len(reader.pages)

        text_parts: list[str] = []
        for i in range(min(max_pages, len(reader.pages))):
            t = reader.pages[i].extract_text() or ""
            if t.strip():
                text_parts.append(t)

        text = _clean_text("\n".join(text_parts))

        if len(text) < 500:
            ocr_text = _ocr_pdf_file(path, max_pages=min(PDF_OCR_MAX_PAGES, info["pages"] or PDF_OCR_MAX_PAGES))
            if len(ocr_text) > len(text):
                text = ocr_text
                info["ocr_used"] = True

        info["text_len"] = len(text)
        text_lc = text.lower()

        hits: dict[str, int] = {}
        for kw in ESTIMATOR_KEYWORDS:
            if kw in text_lc:
                hits[kw] = text_lc.count(kw)
        info["keyword_hits"] = dict(sorted(hits.items(), key=lambda x: x[1], reverse=True)[:25])

        info["date_candidates"] = _extract_date_candidates_from_text(text)
        info["snippet"] = text[:1000]
        info["text"] = text[:40000]
        info["requirements"] = _extract_requirements(text)

    except Exception as e:
        info["error"] = f"PDF read failed: {e}"

    return info


def extract_docx_info(path: Path, max_paras: int = DOCX_MAX_PARAS) -> dict[str, Any]:
    info: dict[str, Any] = {
        "headings": [],
        "keyword_hits": {},
        "snippet": "",
        "text_len": 0,
        "requirements": {"strict": [], "loose": []},
        "date_candidates": [],
    }

    try:
        doc = Document(str(path))

        paras = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        paras = paras[:max_paras]

        table_lines: list[str] = []
        for table in doc.tables[:DOCX_MAX_TABLES]:
            for row in table.rows[:DOCX_MAX_TABLE_ROWS]:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    table_lines.append(" ".join(cells)[:400])

        text = _clean_text("\n".join(paras + table_lines))
        info["text_len"] = len(text)

        headings: list[str] = []
        for p in doc.paragraphs:
            if p.style and p.style.name and "Heading" in p.style.name and p.text.strip():
                headings.append(p.text.strip())
            if len(headings) >= 40:
                break
        info["headings"] = headings

        text_lc = text.lower()
        hits: dict[str, int] = {}
        for kw in ESTIMATOR_KEYWORDS:
            if kw in text_lc:
                hits[kw] = text_lc.count(kw)
        info["keyword_hits"] = dict(sorted(hits.items(), key=lambda x: x[1], reverse=True)[:25])

        info["date_candidates"] = _extract_date_candidates_from_text(text)
        info["snippet"] = text[:1000]
        info["text"] = text[:40000]
        info["requirements"] = _extract_requirements(text)

    except Exception as e:
        info["error"] = f"DOCX read failed: {e}"

    return info


def detect_boq_columns(header_row: list[str]) -> dict[str, int]:
    cols = [c.lower().strip() for c in header_row]

    def find_any(keys: list[str]) -> int:
        for i, c in enumerate(cols):
            if any(k in c for k in keys):
                return i
        return -1

    mapping = {
        "item": find_any(["item", "no", "line", "ref"]),
        "description": find_any(["description", "desc", "details", "work", "item description"]),
        "qty": find_any(["qty", "quantity", "quant"]),
        "unit": find_any(["unit", "uom"]),
        "rate": find_any(["rate", "price"]),
        "amount": find_any(["amount", "total", "value"]),
    }
    return {k: v for k, v in mapping.items() if v != -1}


def extract_xlsx_info(path: Path) -> dict[str, Any]:
    info: dict[str, Any] = {"sheets": []}
    try:
        wb = load_workbook(filename=str(path), data_only=True, read_only=True)
        for name in wb.sheetnames[:XLSX_MAX_SHEETS]:
            ws = wb[name]
            rows: list[list[str]] = []
            for r in ws.iter_rows(min_row=1, max_row=XLSX_MAX_ROWS, values_only=True):
                rows.append([("" if v is None else str(v)).strip() for v in r][:30])

            header = rows[0] if rows else []
            colmap = detect_boq_columns(header) if header else {}

            info["sheets"].append({
                "name": name,
                "header_guess": header[:20],
                "boq_column_map": colmap,
                "preview_rows": rows[:8],
            })
    except Exception as e:
        info["error"] = f"XLSX read failed: {e}"
    return info


def extract_csv_info(path: Path) -> dict[str, Any]:
    info: dict[str, Any] = {"header_guess": [], "boq_column_map": {}, "preview_rows": []}
    try:
        with open(path, "r", encoding="utf-8", errors="ignore", newline="") as f:
            reader = csv.reader(f)
            rows: list[list[str]] = []
            for i, r in enumerate(reader):
                rows.append([c.strip() for c in r][:30])
                if i >= CSV_MAX_ROWS:
                    break

        header = rows[0] if rows else []
        info["header_guess"] = header[:20]
        info["boq_column_map"] = detect_boq_columns(header) if header else {}
        info["preview_rows"] = rows[:8]
    except Exception as e:
        info["error"] = f"CSV read failed: {e}"
    return info


def extract_image_info(path: Path) -> dict[str, Any]:
    info: dict[str, Any] = {
        "text_len": 0,
        "snippet": "",
        "requirements": {"strict": [], "loose": []},
        "date_candidates": [],
        "ocr_used": False,
    }
    txt = _ocr_image_file(path)
    if txt:
        info["ocr_used"] = True
        info["text_len"] = len(txt)
        info["snippet"] = txt[:1000]
        info["text"] = txt[:12000]
        info["date_candidates"] = _extract_date_candidates_from_text(txt)
        info["requirements"] = _extract_requirements(txt)
    return info


def extract_by_type(path: Path, category: str) -> dict[str, Any]:
    ext = path.suffix.lower()

    if ext == ".pdf":
        return extract_pdf_info(path)
    if ext == ".docx":
        return extract_docx_info(path)
    if ext in {".xlsx", ".xls"}:
        return extract_xlsx_info(path)
    if ext == ".csv":
        return extract_csv_info(path)
    if ext in IMAGE_EXTS:
        return extract_image_info(path)

    if category == "drawings":
        return {
            "drawing_number_guess": guess_drawing_number(path.name),
            "revision_guess": guess_revision(path.name),
        }

    return {}


# ---------------- fact extraction ----------------
def _first_clean_match(text: str, patterns: list[str], max_len: int = 180) -> str | None:
    if not text:
        return None

    for pat in patterns:
        for m in re.finditer(pat, text, flags=re.IGNORECASE):
            val = m.group(1) if m.groups() else m.group(0)
            val = _normalize_line(val)
            if not val:
                continue
            if _looks_irrelevant(val) or _looks_like_schedule_row(val) or _is_gibberish_line(val):
                continue
            return val[:max_len]
    return None


def _clean_fact_value(text: str | None, max_len: int = 220) -> str | None:
    if not text:
        return None

    s = _normalize_line(text)
    if not s:
        return None
    if _looks_irrelevant(s) or _looks_like_schedule_row(s) or _is_gibberish_line(s):
        return None

    s = re.split(r"\b(employer'?s agent|project manager|quantity surveyor|name:)\b", s, maxsplit=1, flags=re.I)[0].strip(" :;,-")
    s = re.split(r"\b(section|clause|appendix|schedule)\b\s+\d", s, maxsplit=1, flags=re.I)[0].strip(" :;,-")
    s = re.sub(r"\s+", " ", s).strip()

    if len(s.split()) < 2:
        return None
    return s[:max_len]


def _extract_project_address(text: str) -> str | None:
    if not text:
        return None

    patterns = [
        r"\b(?:site address|project address|works address|property address)\b\s*[:\-]?\s*([^\n]{10,180})",
        r"\b(?:located at|situated at)\b\s*([^\n]{10,160})",
    ]
    raw = _first_clean_match(text, patterns, 180)
    if raw:
        raw = _clean_fact_value(raw, 180)
        if raw and "email:" not in raw.lower() and "@" not in raw:
            return raw

    postcode = re.search(r"\b([A-Z]{1,2}\d[A-Z\d]?\s*\d[A-Z]{2})\b", text, flags=re.I)
    if postcode:
        idx = postcode.start()
        start = max(0, idx - 80)
        snippet = _normalize_line(text[start: postcode.end()])
        if 6 <= len(snippet.split()) <= 24 and "@" not in snippet:
            return snippet[:180]

    return None


def _extract_client_name(text: str) -> str | None:
    patterns = [
        r"\b(?:client|employer|contract administrator)\b\s*[:\-]?\s*([^\n\.]{4,120})",
    ]
    return _clean_fact_value(_first_clean_match(text, patterns, 120), 120)


def _extract_contract_type(text: str) -> str | None:
    patterns = [
        r"\b(JCT\s*[A-Z0-9 \-/]{0,40})\b",
        r"\b(NEC\s*[A-Z0-9 \-/]{0,40})\b",
        r"\b(IFC|IChemE|FIDIC)\b",
    ]
    return _clean_fact_value(_first_clean_match(text, patterns, 120), 120)


def _extract_submission_route(text: str) -> str | None:
    if not text:
        return None

    email = EMAIL_RE.search(text)
    if email:
        return f"Email: {email.group(1)}"

    if re.search(r"\bportal\b", text, flags=re.I):
        return "Portal submission"

    if re.search(r"\bupload\b", text, flags=re.I):
        return "Upload submission"

    if re.search(r"\be-tender\b|\beprocurement\b", text, flags=re.I):
        return "Electronic tender portal"

    return None


def _extract_deadline_value(text: str, fallback_dates: list[str]) -> str | None:
    if text:
        patterns = [
            r"\b(?:tender|return|submission|submit)\s+(?:date|deadline|by)\b\s*[:\-]?\s*([^\n\.]{4,80})",
            r"\b(?:closing date|deadline|tender return)\b\s*[:\-]?\s*([^\n\.]{4,80})",
        ]
        raw = _first_clean_match(text, patterns, 100)
        if raw:
            date_inside = re.search(
                r"\b(\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\s+\d{2,4}|\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4})\b",
                raw,
                flags=re.I,
            )
            if date_inside and _valid_date_candidate(date_inside.group(1)):
                return date_inside.group(1)

    clean_dates = [d for d in fallback_dates if _valid_date_candidate(d)]
    return clean_dates[0] if clean_dates else None


def _extract_start_date(text: str) -> str | None:
    patterns = [
        r"\b(?:start date|commencement date|commence on|works start)\b\s*[:\-]?\s*([^\n\.]{4,80})",
    ]
    raw = _first_clean_match(text, patterns, 100)
    if raw:
        m = re.search(r"\b(\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\s+\d{2,4}|\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4})\b", raw, re.I)
        if m and _valid_date_candidate(m.group(1)):
            return m.group(1)
    return None


def _extract_completion_date(text: str) -> str | None:
    patterns = [
        r"\b(?:completion date|date for completion|complete by|practical completion)\b\s*[:\-]?\s*([^\n\.]{4,80})",
    ]
    raw = _first_clean_match(text, patterns, 100)
    if raw:
        m = re.search(r"\b(\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\s+\d{2,4}|\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4})\b", raw, re.I)
        if m and _valid_date_candidate(m.group(1)):
            return m.group(1)
    return None


def _extract_programme_value(text: str) -> str | None:
    if not text:
        return None

    patterns = [
        r"\b(?:programme|program|duration|contract period)\b\s*[:\-]?\s*(\d{1,3}\s*(?:weeks?|months?))",
        r"(\d{1,3}\s*(?:weeks?|months?))\s*\b(?:programme|duration|contract period)\b",
    ]
    return _first_clean_match(text, patterns, 80)


def _extract_working_hours_value(text: str) -> str | None:
    if not text:
        return None

    patterns = [
        r"\b(?:working hours|site hours|hours of work)\b\s*[:\-]?\s*([^\n\.]{6,120})",
        r"\b((?:mon|monday|tue|tuesday|wed|wednesday|thu|thursday|fri|friday)[^\n\.]{0,80}\d{1,2}[:.]\d{2}[^\n\.]{0,20}\d{1,2}[:.]\d{2})\b",
    ]

    raw = _first_clean_match(text, patterns, 140)
    if not raw:
        return None

    bad_phrases = [
        "scaffold",
        "secured outside normal working hours",
        "trespass",
        "ladder",
        "safety",
        "where possible",
    ]
    if any(x in raw.lower() for x in bad_phrases):
        return None

    return raw


def _extract_retention_value(text: str) -> str | None:
    if not text:
        return None

    patterns = [
        r"\bretention\b\s*[:\-]?\s*(\d{1,2}(?:\.\d+)?\s*%)",
        r"(\d{1,2}(?:\.\d+)?\s*%)\s*\bretention\b",
        r"\bretention\b.{0,40}\b(of\s+\d{1,2}(?:\.\d+)?\s*%)",
    ]
    return _first_clean_match(text, patterns, 80)


def _extract_ld_value(text: str) -> str | None:
    if not text:
        return None

    patterns = [
        r"\b(?:liquidated damages|LDs?|LADs?)\b.{0,50}(£\s?\d[\d,]*\.?\d*(?:\s*(?:per week|per day|per month))?)",
        r"(£\s?\d[\d,]*\.?\d*\s*(?:per week|per day|per month))",
    ]
    return _first_clean_match(text, patterns, 120)


def _extract_insurance_value(text: str) -> str | None:
    if not text:
        return None

    patterns = [
        r"\b(?:public liability|employers liability|EL|PL)\b.{0,50}(£\s?\d[\d,]*\.?\d*)",
        r"\binsurance\b.{0,50}(£\s?\d[\d,]*\.?\d*)",
    ]
    return _first_clean_match(text, patterns, 120)


def _extract_rectification_period(text: str) -> str | None:
    patterns = [
        r"\b(?:rectification period|defects liability period)\b\s*[:\-]?\s*([^\n\.]{4,80})",
        r"\b((?:\d{1,2}\s*(?:months?|weeks?)))\b.{0,40}\b(?:rectification period|defects liability)\b",
    ]
    return _clean_fact_value(_first_clean_match(text, patterns, 100), 100)


def _extract_bond_warranty(text: str) -> str | None:
    patterns = [
        r"\b(?:bond|performance bond|parent company guarantee|warranty|collateral warranty)\b\s*[:\-]?\s*([^\n\.]{4,120})",
    ]
    return _clean_fact_value(_first_clean_match(text, patterns, 120), 120)


def _extract_accreditations(text: str) -> list[str]:
    found: list[str] = []
    checks = ["CHAS", "SMAS", "SafeContractor", "Constructionline", "ISO 9001", "ISO 14001", "ISO 45001"]
    for c in checks:
        if re.search(rf"\b{re.escape(c)}\b", text or "", flags=re.I):
            found.append(c)
    return _trim_list(found, 6, 80)


def _extract_fact_sentence(text: str, keywords: list[str]) -> str | None:
    for sentence in _split_sentences(text):
        s = _clean_candidate_sentence(sentence, 220)
        if not s:
            continue
        sl = s.lower()
        if any(k in sl for k in keywords):
            return s
    return None


def _find_bucket_evidence(merged: str, needle: str, bucket: str, max_items: int = 1) -> list[str]:
    merged_lc = merged.lower()
    out: list[str] = []
    start = 0

    while len(out) < max_items:
        idx = merged_lc.find(needle, start)
        if idx == -1:
            break

        s = _sentences_around(merged, idx, max_sentences=2)
        s = _normalize_line(s)[:440]

        if s and not _looks_irrelevant(s) and not _looks_like_schedule_row(s) and not _is_gibberish_line(s):
            sl = s.lower()

            if bucket == "Services / isolations":
                if not any(x in sl for x in ["isolation", "isolations", "live", "disconnect", "disconnection", "divert", "diversion"]):
                    start = idx + len(needle)
                    continue

            if _sentence_score(s) > 0 and s not in out:
                out.append(s)

        start = idx + len(needle)

    return out


def _shorten_constraint(bucket: str, text: str) -> str | None:
    s = _clean_fact_value(text, 260)
    if not s:
        return None

    bucket_prefix = {
        "Asbestos / hazardous materials": "Asbestos or hazardous materials appear to be referenced.",
        "Temporary works / propping": "Temporary works or propping requirements appear to be referenced.",
        "Party wall / adjacent structures": "Adjacent structure or third-party interface risk appears to be referenced.",
        "Traffic management / access": "Access or logistics constraints appear to be referenced.",
        "Noise / dust / vibration": "Noise, dust, or vibration controls appear to be referenced.",
        "Services / isolations": "Live services or isolation requirements appear to be referenced.",
        "Waste / crushing / segregation": "Waste handling or segregation requirements appear to be referenced.",
        "Permits / licences": "Permits, licences, or consents appear to be referenced.",
        "Working hours / constraints": "Working hour constraints appear to be referenced.",
        "Liquidated damages / penalties": "Delay damages or penalties appear to be referenced.",
    }

    generic = bucket_prefix.get(bucket, "")
    if len(s.split()) < 9 or BAD_FRAGMENT_RE.search(s):
        return generic or None
    return s[:260]


# ---------------- briefing ----------------
def extract_pack_briefing(sections: dict[str, list[dict[str, Any]]]) -> dict[str, Any]:
    text_blobs: list[str] = []
    strict_reqs: list[str] = []
    loose_reqs: list[str] = []

    for cat in ["forms", "prelims", "addenda", "specs", "documents", "pdfs", "photos"]:
        for item in sections.get(cat, []):
            ex = item.get("extracted") or {}
            t = ex.get("text")
            if t and isinstance(t, str) and t.strip():
                if _looks_irrelevant(t[:900]):
                    continue
                text_blobs.append(t)

            req = ex.get("requirements") or {}
            strict_reqs.extend(req.get("strict", []) or [])
            loose_reqs.extend(req.get("loose", []) or [])

    merged = "\n\n".join(text_blobs)[:250000]
    merged_lc = merged.lower()

    date_candidates: set[str] = set()
    for _, items in sections.items():
        for it in items:
            ex = it.get("extracted") or {}
            for d in ex.get("date_candidates", []) or []:
                if _valid_date_candidate(str(d)):
                    date_candidates.add(str(d))

    all_dates = sorted(date_candidates)[:40]

    site_address = _extract_project_address(merged)
    headline_client = _extract_client_name(merged)
    headline_contract = _extract_contract_type(merged)
    headline_deadline = _extract_deadline_value(merged, all_dates)
    headline_submission = _extract_submission_route(merged)
    headline_start = _extract_start_date(merged)
    headline_completion = _extract_completion_date(merged)
    headline_prog = _extract_programme_value(merged)
    headline_hours = _extract_working_hours_value(merged)
    headline_ret = _extract_retention_value(merged)
    headline_ld = _extract_ld_value(merged)
    headline_ins = _extract_insurance_value(merged)
    headline_rectification = _extract_rectification_period(merged)
    headline_bond = _extract_bond_warranty(merged)
    accreditations = _extract_accreditations(merged)

    headline_asbestos = _extract_fact_sentence(merged, ["asbestos", "acm", "hazardous"])
    headline_services = _extract_fact_sentence(merged, ["live services", "isolation", "disconnect", "diversion"])
    headline_access = _extract_fact_sentence(merged, ["access", "logistics", "delivery", "traffic management"])
    headline_permits = _extract_fact_sentence(merged, ["permit", "permits", "licence", "license", "consent"])

    strict_clean = _dedup_keep_best(strict_reqs, 18, 340)
    loose_clean = _dedup_keep_best(loose_reqs, 14, 340)

    constraints: list[str] = []
    for bucket, needles in RISK_BUCKETS.items():
        if not any(n in merged_lc for n in needles):
            continue

        best_ev = ""
        for needle in needles:
            if needle in merged_lc:
                evs = _find_bucket_evidence(merged, needle, bucket=bucket, max_items=1)
                if evs:
                    best_ev = evs[0]
                    break

        if best_ev:
            short = _shorten_constraint(bucket, best_ev)
            if short and short not in constraints:
                constraints.append(short)

        if len(constraints) >= 10:
            break

    pricing_watchouts = strict_clean[:8] if strict_clean else constraints[:8]

    missing: list[str] = []
    if not headline_deadline and not all_dates:
        missing.append("Tender return date / deadline not found")
    if not headline_submission:
        missing.append("Submission method (email/portal/upload) not found")
    if not headline_prog and not headline_start and not headline_completion:
        missing.append("Programme / duration not found")
    if not headline_hours:
        missing.append("Working hours not found")
    if not headline_ret:
        missing.append("Retention not found")
    if not headline_ld:
        missing.append("Liquidated damages / LADs not found")
    if not headline_ins:
        missing.append("Insurance levels not found")
    if not accreditations:
        missing.append("Accreditations (CHAS/SMAS/etc) not found")

    clarifications = missing[:]
    if not headline_submission and "Submission method (email/portal/upload) not found" not in clarifications:
        clarifications.append("Confirm whether the tender must be emailed, uploaded, or submitted through a portal.")
    if not headline_prog and "Programme / duration not found" not in clarifications:
        clarifications.append("Confirm programme duration, milestones, and phased handover requirements.")
    clarifications = clarifications[:10]

    executive_lines: list[str] = []
    if site_address:
        executive_lines.append(f"Project / site address: {site_address}")
    if headline_client:
        executive_lines.append(f"Client / employer: {headline_client}")
    if headline_contract:
        executive_lines.append(f"Contract type: {headline_contract}")
    if headline_deadline:
        executive_lines.append(f"Deadline / key date: {headline_deadline}")
    if headline_submission:
        executive_lines.append(f"Submission route: {headline_submission}")
    if headline_start:
        executive_lines.append(f"Start date: {headline_start}")
    if headline_completion:
        executive_lines.append(f"Completion date: {headline_completion}")
    if headline_prog:
        executive_lines.append(f"Programme / duration: {headline_prog}")
    if headline_hours:
        executive_lines.append(f"Working hours: {headline_hours}")
    if headline_ret:
        executive_lines.append(f"Retention: {headline_ret}")
    if headline_ld:
        executive_lines.append(f"LD / LADs: {headline_ld}")
    if headline_ins:
        executive_lines.append(f"Insurance: {headline_ins}")
    if headline_rectification:
        executive_lines.append(f"Rectification / defects period: {headline_rectification}")
    if headline_bond:
        executive_lines.append(f"Bond / warranty: {headline_bond}")
    if accreditations:
        executive_lines.append(f"Accreditations: {', '.join(accreditations)}")

    if not executive_lines:
        executive_lines.append("No clear commercial headline terms were confidently extracted from the scanned pack.")

    overview_sentences: list[str] = []
    if site_address:
        overview_sentences.append(f"The works appear to relate to the site at {site_address}.")
    if headline_client:
        overview_sentences.append(f"The client or employer appears to be {headline_client}.")
    if headline_deadline:
        overview_sentences.append(f"The tender appears to close on {headline_deadline}.")
    if headline_submission:
        overview_sentences.append(f"The submission route appears to be {headline_submission.lower()}.")
    if headline_prog:
        overview_sentences.append(f"Programme information indicates {headline_prog.lower()}.")
    if not overview_sentences:
        overview_sentences.append("The pack has been scanned and the strongest commercial, submission, and delivery points are summarised below.")

    commercial_summary_parts: list[str] = []
    if headline_ld:
        commercial_summary_parts.append(f"Liquidated damages appear to be referenced as {headline_ld}.")
    if headline_ret:
        commercial_summary_parts.append(f"Retention appears to be referenced as {headline_ret}.")
    if headline_ins:
        commercial_summary_parts.append(f"Insurance requirements appear to reference {headline_ins}.")
    if headline_bond:
        commercial_summary_parts.append(f"Security or warranty requirements appear to reference {headline_bond}.")
    if not commercial_summary_parts:
        commercial_summary_parts.append("No strong commercial headline terms were confidently extracted beyond the core tender instructions.")

    programme_access_parts: list[str] = []
    if headline_prog:
        programme_access_parts.append(f"Programme duration appears to be {headline_prog}.")
    if headline_start:
        programme_access_parts.append(f"Start date appears to be {headline_start}.")
    if headline_completion:
        programme_access_parts.append(f"Completion date appears to be {headline_completion}.")
    if headline_hours:
        programme_access_parts.append(f"Working hour restrictions appear to be {headline_hours}.")
    if headline_access:
        programme_access_parts.append(headline_access)
    if not programme_access_parts and constraints:
        programme_access_parts.extend(constraints[:2])
    if not programme_access_parts:
        programme_access_parts.append("No strong programme or access constraints were confidently extracted.")

    risk_parts: list[str] = []
    if headline_asbestos:
        risk_parts.append(headline_asbestos)
    if headline_services:
        risk_parts.append(headline_services)
    if headline_permits:
        risk_parts.append(headline_permits)
    risk_parts.extend([c for c in constraints if c not in risk_parts][:4])
    if not risk_parts:
        risk_parts.append("No major delivery risks were confidently extracted from the scanned text.")

    notable_items: list[str] = []
    for x in [headline_asbestos, headline_services, headline_access, headline_permits]:
        if x and x not in notable_items:
            notable_items.append(x)
    notable_items.extend([c for c in constraints if c not in notable_items][:6])
    notable_items = notable_items[:10]

    submission_summary = headline_submission or "Submission method was not confidently extracted."
    if headline_deadline:
        submission_summary += f" Tender deadline appears to be {headline_deadline}."

    return {
        "title": "Tender Pack Summary",
        "executive_summary": "EXECUTIVE SUMMARY\n" + "\n".join([f"• {x}" for x in executive_lines]),
        "overview": " ".join(overview_sentences),
        "commercial_summary": " ".join(commercial_summary_parts),
        "programme_access_summary": " ".join(programme_access_parts),
        "risk_summary": " ".join(risk_parts),
        "submission_summary": submission_summary,
        "pricing_watchouts": pricing_watchouts,
        "notable_items": notable_items,
        "clarifications": clarifications,
        "key_facts": {
            "project_address": site_address,
            "client_name": headline_client,
            "contract_type": headline_contract,
            "deadline_or_key_date": headline_deadline,
            "submission_route": headline_submission,
            "start_date": headline_start,
            "completion_date": headline_completion,
            "programme_duration": headline_prog,
            "working_hours": headline_hours,
            "retention": headline_ret,
            "liquidated_damages": headline_ld,
            "insurance_levels": headline_ins,
            "rectification_period": headline_rectification,
            "bond_or_warranty": headline_bond,
            "accreditations": accreditations,
        },
        "dates_found": all_dates,
        "constraints": constraints,
        "requirements_strict": strict_clean,
        "requirements_loose": loose_clean,
        "missing": missing,
        "sources_scanned": len(text_blobs),
        "evidence": {
            "deadline_candidates": [headline_deadline] if headline_deadline else [],
            "submission_candidates": [headline_submission] if headline_submission else [],
            "programme_candidates": [headline_prog] if headline_prog else [],
            "working_hours_candidates": [headline_hours] if headline_hours else [],
            "retention_candidates": [headline_ret] if headline_ret else [],
            "ld_candidates": [headline_ld] if headline_ld else [],
            "insurance_candidates": [headline_ins] if headline_ins else [],
            "accreditations_candidates": accreditations[:3],
        },
    }


# ---------------- AI rewrite ----------------
def _build_ai_payload(briefing: dict[str, Any]) -> dict[str, Any]:
    return {
        "overview": briefing.get("overview", ""),
        "commercial_summary": briefing.get("commercial_summary", ""),
        "programme_access_summary": briefing.get("programme_access_summary", ""),
        "risk_summary": briefing.get("risk_summary", ""),
        "submission_summary": briefing.get("submission_summary", ""),
        "key_facts": briefing.get("key_facts", {}),
        "dates_found": briefing.get("dates_found", [])[:12],
        "constraints": briefing.get("constraints", [])[:8],
        "requirements_strict": briefing.get("requirements_strict", [])[:12],
        "requirements_loose": briefing.get("requirements_loose", [])[:10],
        "pricing_watchouts": briefing.get("pricing_watchouts", [])[:10],
        "notable_items": briefing.get("notable_items", [])[:10],
        "clarifications": briefing.get("clarifications", [])[:10],
        "missing": briefing.get("missing", [])[:10],
    }


def ai_enhance_briefing(briefing: dict[str, Any]) -> dict[str, Any]:
    if not openai_client:
        briefing["ai_used"] = False
        briefing["ai_error"] = "OPENAI_API_KEY not set"
        return briefing

    payload = _build_ai_payload(briefing)

    system_prompt = """
You are a senior UK construction estimator reviewing a tender pack.

You will receive already-curated tender facts.
Rewrite them into a clean, practical tender summary page that reads like a professional report.

Rules:
- Use ONLY the supplied facts.
- Do NOT invent missing information.
- Do NOT mention file names, extraction, OCR, evidence, or sources.
- Keep the wording commercially useful and easy to scan.
- Do not combine unrelated facts into one sentence.
- If something is unclear, leave it out rather than forcing it.
- Avoid vague filler phrases.
- Keep the summary polished, concise, and estimator-friendly.

Return STRICT JSON with this exact structure:
{
  "title": "string",
  "executive_summary": "string",
  "overview": "string",
  "commercial_summary": "string",
  "programme_access_summary": "string",
  "risk_summary": "string",
  "submission_summary": "string",
  "pricing_watchouts": ["..."],
  "notable_items": ["..."],
  "clarifications": ["..."],
  "key_facts": {
    "project_address": "string or null",
    "client_name": "string or null",
    "contract_type": "string or null",
    "deadline_or_key_date": "string or null",
    "submission_route": "string or null",
    "start_date": "string or null",
    "completion_date": "string or null",
    "programme_duration": "string or null",
    "working_hours": "string or null",
    "retention": "string or null",
    "liquidated_damages": "string or null",
    "insurance_levels": "string or null",
    "rectification_period": "string or null",
    "bond_or_warranty": "string or null",
    "accreditations": ["..."]
  },
  "constraints": ["..."],
  "requirements_strict": ["..."],
  "requirements_loose": ["..."],
  "missing": ["..."]
}
"""

    try:
        resp = openai_client.chat.completions.create(
            model=OPENAI_MODEL,
            temperature=0.1,
            timeout=OPENAI_TIMEOUT_SEC,
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": json.dumps(payload, ensure_ascii=False)},
            ],
        )

        content = resp.choices[0].message.content or ""
        parsed = _safe_json_loads(content)
        if not parsed:
            briefing["ai_used"] = False
            briefing["ai_error"] = "OpenAI returned non-JSON content"
            return briefing

        enhanced = {
            "title": parsed.get("title") or briefing.get("title") or "Tender Pack Summary",
            "executive_summary": parsed.get("executive_summary") or briefing.get("executive_summary"),
            "overview": parsed.get("overview") or briefing.get("overview"),
            "commercial_summary": parsed.get("commercial_summary") or briefing.get("commercial_summary"),
            "programme_access_summary": parsed.get("programme_access_summary") or briefing.get("programme_access_summary"),
            "risk_summary": parsed.get("risk_summary") or briefing.get("risk_summary"),
            "submission_summary": parsed.get("submission_summary") or briefing.get("submission_summary"),
            "pricing_watchouts": _trim_list(parsed.get("pricing_watchouts"), 10, 340) or briefing.get("pricing_watchouts", []),
            "notable_items": _trim_list(parsed.get("notable_items"), 10, 340) or briefing.get("notable_items", []),
            "clarifications": _trim_list(parsed.get("clarifications"), 10, 340) or briefing.get("clarifications", []),
            "key_facts": {
                "project_address": (parsed.get("key_facts") or {}).get("project_address") or (briefing.get("key_facts") or {}).get("project_address"),
                "client_name": (parsed.get("key_facts") or {}).get("client_name") or (briefing.get("key_facts") or {}).get("client_name"),
                "contract_type": (parsed.get("key_facts") or {}).get("contract_type") or (briefing.get("key_facts") or {}).get("contract_type"),
                "deadline_or_key_date": (parsed.get("key_facts") or {}).get("deadline_or_key_date") or (briefing.get("key_facts") or {}).get("deadline_or_key_date"),
                "submission_route": (parsed.get("key_facts") or {}).get("submission_route") or (briefing.get("key_facts") or {}).get("submission_route"),
                "start_date": (parsed.get("key_facts") or {}).get("start_date") or (briefing.get("key_facts") or {}).get("start_date"),
                "completion_date": (parsed.get("key_facts") or {}).get("completion_date") or (briefing.get("key_facts") or {}).get("completion_date"),
                "programme_duration": (parsed.get("key_facts") or {}).get("programme_duration") or (briefing.get("key_facts") or {}).get("programme_duration"),
                "working_hours": (parsed.get("key_facts") or {}).get("working_hours") or (briefing.get("key_facts") or {}).get("working_hours"),
                "retention": (parsed.get("key_facts") or {}).get("retention") or (briefing.get("key_facts") or {}).get("retention"),
                "liquidated_damages": (parsed.get("key_facts") or {}).get("liquidated_damages") or (briefing.get("key_facts") or {}).get("liquidated_damages"),
                "insurance_levels": (parsed.get("key_facts") or {}).get("insurance_levels") or (briefing.get("key_facts") or {}).get("insurance_levels"),
                "rectification_period": (parsed.get("key_facts") or {}).get("rectification_period") or (briefing.get("key_facts") or {}).get("rectification_period"),
                "bond_or_warranty": (parsed.get("key_facts") or {}).get("bond_or_warranty") or (briefing.get("key_facts") or {}).get("bond_or_warranty"),
                "accreditations": _trim_list((parsed.get("key_facts") or {}).get("accreditations"), 6, 80) or (briefing.get("key_facts") or {}).get("accreditations", []),
            },
            "dates_found": briefing.get("dates_found", []),
            "constraints": _trim_list(parsed.get("constraints"), 10, 340) or briefing.get("constraints", []),
            "requirements_strict": _trim_list(parsed.get("requirements_strict"), 14, 340) or briefing.get("requirements_strict", []),
            "requirements_loose": _trim_list(parsed.get("requirements_loose"), 12, 340) or briefing.get("requirements_loose", []),
            "missing": _trim_list(parsed.get("missing"), 10, 220) or briefing.get("missing", []),
            "sources_scanned": briefing.get("sources_scanned", 0),
            "evidence": briefing.get("evidence", {}),
            "ai_used": True,
            "ai_error": None,
        }
        return enhanced

    except Exception as e:
        briefing["ai_used"] = False
        briefing["ai_error"] = str(e)
        return briefing


# ---------------- PDF export ----------------
def build_pdf_report(report: dict[str, Any]) -> bytes:
    briefing = report.get("briefing") or {}
    summary = report.get("summary") or {}
    key_facts = briefing.get("key_facts") or {}

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        rightMargin=16 * mm,
        leftMargin=16 * mm,
        topMargin=14 * mm,
        bottomMargin=14 * mm,
        title=(briefing.get("title") or "Tender Pack Summary"),
    )

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="TitleSmall", parent=styles["Title"], fontSize=20, leading=24, textColor=colors.HexColor("#111827")))
    styles.add(ParagraphStyle(name="SectionHead", parent=styles["Heading2"], fontSize=12, leading=15, textColor=colors.HexColor("#374151"), spaceAfter=6))
    styles.add(ParagraphStyle(name="BodySmall", parent=styles["BodyText"], fontSize=9.5, leading=13, textColor=colors.HexColor("#111827")))
    styles.add(ParagraphStyle(name="Meta", parent=styles["BodyText"], fontSize=8.5, leading=11, textColor=colors.HexColor("#6B7280")))

    story: list[Any] = []

    def esc(s: Any) -> str:
        return str(s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    def add_paragraph(text: str, style: str = "BodySmall") -> None:
        if text:
            story.append(Paragraph(esc(text), styles[style]))
            story.append(Spacer(1, 3 * mm))

    def add_bullets(items: list[str]) -> None:
        cleaned = [i for i in items if isinstance(i, str) and i.strip()]
        if not cleaned:
            return
        flow = ListFlowable(
            [ListItem(Paragraph(esc(i), styles["BodySmall"])) for i in cleaned],
            bulletType="bullet",
            leftPadding=14,
        )
        story.append(flow)
        story.append(Spacer(1, 3 * mm))

    title = briefing.get("title") or "Tender Pack Summary"
    story.append(Paragraph(esc(title), styles["TitleSmall"]))
    story.append(
        Paragraph(
            esc(f"Generated {datetime.utcnow().strftime('%d %b %Y %H:%M UTC')} • Files scanned: {summary.get('total_files_scanned', 0)}"),
            styles["Meta"],
        )
    )
    story.append(Spacer(1, 4 * mm))

    story.append(Paragraph("Executive Summary", styles["SectionHead"]))
    add_paragraph(briefing.get("overview") or briefing.get("executive_summary") or "")

    story.append(Paragraph("Commercial Summary", styles["SectionHead"]))
    add_paragraph(briefing.get("commercial_summary") or "")

    story.append(Paragraph("Programme and Access", styles["SectionHead"]))
    add_paragraph(briefing.get("programme_access_summary") or "")

    story.append(Paragraph("Risk Summary", styles["SectionHead"]))
    add_paragraph(briefing.get("risk_summary") or "")

    story.append(Paragraph("Submission Summary", styles["SectionHead"]))
    add_paragraph(briefing.get("submission_summary") or "")

    story.append(Paragraph("Key Facts", styles["SectionHead"]))
    facts: list[str] = []
    fact_map = [
        ("Project / site address", key_facts.get("project_address")),
        ("Client / employer", key_facts.get("client_name")),
        ("Contract type", key_facts.get("contract_type")),
        ("Deadline / key date", key_facts.get("deadline_or_key_date")),
        ("Submission route", key_facts.get("submission_route")),
        ("Start date", key_facts.get("start_date")),
        ("Completion date", key_facts.get("completion_date")),
        ("Programme / duration", key_facts.get("programme_duration")),
        ("Working hours", key_facts.get("working_hours")),
        ("Retention", key_facts.get("retention")),
        ("LD / LADs", key_facts.get("liquidated_damages")),
        ("Insurance", key_facts.get("insurance_levels")),
        ("Rectification period", key_facts.get("rectification_period")),
        ("Bond / warranty", key_facts.get("bond_or_warranty")),
    ]
    for label, value in fact_map:
        if value:
            facts.append(f"{label}: {value}")
    if key_facts.get("accreditations"):
        facts.append("Accreditations: " + ", ".join(key_facts["accreditations"][:6]))
    add_bullets(facts or ["No key facts were confidently extracted."])

    story.append(Paragraph("Pricing Watchouts", styles["SectionHead"]))
    add_bullets(briefing.get("pricing_watchouts") or ["No pricing watchouts were confidently extracted."])

    story.append(Paragraph("Notable Items", styles["SectionHead"]))
    add_bullets(briefing.get("notable_items") or ["No notable items were confidently extracted."])

    story.append(Paragraph("Key Constraints", styles["SectionHead"]))
    add_bullets(briefing.get("constraints") or ["No strong constraints were confidently extracted."])

    story.append(Paragraph("Mandatory Requirements", styles["SectionHead"]))
    add_bullets((briefing.get("requirements_strict") or [])[:12] or ["No strong mandatory requirements were confidently extracted."])

    story.append(Paragraph("Other Submission Requests", styles["SectionHead"]))
    add_bullets((briefing.get("requirements_loose") or [])[:10] or ["No other clear submission requests were confidently extracted."])

    story.append(Paragraph("Clarifications / Queries to Raise", styles["SectionHead"]))
    add_bullets(briefing.get("clarifications") or briefing.get("missing") or ["No additional clarifications suggested."])

    doc.build(story)
    return buf.getvalue()


def _strip_internal_fields(sections: dict[str, list[dict[str, Any]]]) -> None:
    for items in sections.values():
        for it in items:
            ex = it.get("extracted")
            if isinstance(ex, dict):
                ex.pop("text", None)
                if isinstance(ex.get("sheets"), list):
                    for sh in ex["sheets"]:
                        if isinstance(sh, dict):
                            sh.pop("preview_rows", None)


# ---------------- routes ----------------
@app.get("/", response_class=HTMLResponse)
def home(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})


async def _save_upload_to_path(uf: UploadFile, dest: Path, max_bytes: int) -> int:
    dest.parent.mkdir(parents=True, exist_ok=True)
    size = 0
    with open(dest, "wb") as f:
        while True:
            chunk = await uf.read(1024 * 1024)
            if not chunk:
                break
            size += len(chunk)
            if size > max_bytes:
                raise ValueError("File too large")
            f.write(chunk)
    return size


@app.post("/api/report/pdf")
async def report_pdf(payload: dict[str, Any] = Body(...)):
    pdf_bytes = build_pdf_report(payload)
    filename = "tender-summary.pdf"
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.post("/api/analyse")
async def analyse(
    zip_file: Optional[list[UploadFile]] = File(None),
    files: Optional[list[UploadFile]] = File(None),
    file: Optional[UploadFile] = File(None),
):
    try:
        uploads: list[UploadFile] = []
        if zip_file:
            uploads.extend(zip_file)
        if files:
            uploads.extend(files)
        if file:
            uploads.append(file)

        if not uploads:
            return JSONResponse({"error": "No files uploaded."}, status_code=400)

        sections: dict[str, list[dict[str, Any]]] = defaultdict(list)
        ext_counter: Counter[str] = Counter()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            extract_root = tmp_path / "extracted"
            extract_root.mkdir(parents=True, exist_ok=True)

            uploaded_names: list[str] = []
            scanned_paths: list[tuple[Path, str]] = []

            for uf in uploads:
                if not uf.filename or not _is_supported_upload(uf.filename):
                    continue

                uploaded_names.append(uf.filename)

                if uf.filename.lower().endswith(".zip"):
                    zp = tmp_path / f"upload_{len(uploaded_names)}.zip"
                    try:
                        await _save_upload_to_path(uf, zp, max_bytes=MAX_FILE_BYTES)
                    except ValueError:
                        return JSONResponse({"error": f"File too large (limit 300MB): {uf.filename}"}, status_code=400)

                    try:
                        extracted = _recursive_extract_zip(zp, extract_root, depth=0)
                    except Exception as e:
                        return JSONResponse({"error": f"Could not extract ZIP {uf.filename}: {e}"}, status_code=400)

                    for p in extracted:
                        scanned_paths.append((p, str(p.relative_to(extract_root)).replace("\\", "/")))
                else:
                    rel = (uf.filename or "").replace("\\", "/").lstrip("/")
                    rel_path = Path(rel)
                    safe_rel = Path(*[p for p in rel_path.parts if p not in ("", ".", "..")])
                    if str(safe_rel) in ("", "."):
                        safe_rel = Path(Path(uf.filename).name)

                    if not _is_supported_upload(str(safe_rel)):
                        continue

                    op = tmp_path / safe_rel
                    try:
                        await _save_upload_to_path(uf, op, max_bytes=MAX_FILE_BYTES)
                    except ValueError:
                        return JSONResponse({"error": f"File too large (limit 300MB): {uf.filename}"}, status_code=400)

                    scanned_paths.append((op, str(safe_rel).replace("\\", "/")))

            total_available_files = len(scanned_paths)
            if total_available_files == 0:
                return JSONResponse({"error": "No supported files found in upload."}, status_code=400)

            total_files = 0
            by_category_count: dict[str, int] = defaultdict(int)

            classified_rows: list[tuple[int, Path, str, str, str]] = []
            for p, display in scanned_paths[:MAX_FILES_TO_PROCESS]:
                total_files += 1
                ext = p.suffix.lower() if p.suffix else "(no_ext)"
                ext_counter[ext] += 1

                category = classify_file(p, display=display)
                by_category_count[category] += 1
                priority = _file_priority_score(display, category)
                classified_rows.append((priority, p, display, ext, category))

            classified_rows.sort(key=lambda x: x[0], reverse=True)

            extraction_jobs: list[tuple[dict[str, Any], Path, str]] = []
            for priority, p, display, ext, category in classified_rows:
                item = {
                    "file": display,
                    "ext": ext,
                    "category": category,
                    "priority": priority,
                    "extracted": {},
                }
                sections[category].append(item)
                extraction_jobs.append((item, p, category))

            def _run_extract(job: tuple[dict[str, Any], Path, str]) -> tuple[dict[str, Any], dict[str, Any]]:
                item, p, category = job
                return item, extract_by_type(p, category)

            if extraction_jobs:
                with ThreadPoolExecutor(max_workers=EXTRACT_WORKERS) as executor:
                    futures = [executor.submit(_run_extract, job) for job in extraction_jobs]
                    for future in as_completed(futures):
                        try:
                            item, extracted = future.result()
                            item["extracted"] = extracted
                        except Exception as e:
                            print(f"EXTRACT ERROR: {e}")

            def has_cat(cat: str) -> bool:
                return len(sections.get(cat, [])) > 0

            raw_briefing = extract_pack_briefing(sections)
            briefing = ai_enhance_briefing(raw_briefing)

            _strip_internal_fields(sections)

            report = {
                "summary": {
                    "uploaded_items": len(uploaded_names),
                    "uploaded_names": uploaded_names[:200],
                    "total_files_scanned": total_files,
                    "total_files_available": total_available_files,
                    "deep_scanned_files": len(extraction_jobs),
                    "by_extension": dict(ext_counter.most_common()),
                    "by_category": dict(sorted(by_category_count.items(), key=lambda x: x[1], reverse=True)),
                    "boq_found": has_cat("boq"),
                    "register_found": has_cat("registers"),
                    "drawings_found": has_cat("drawings"),
                    "forms_found": has_cat("forms"),
                    "prelims_found": has_cat("prelims"),
                    "specs_found": has_cat("specs"),
                    "addenda_found": has_cat("addenda"),
                    "ocr_available": bool(convert_from_path and pytesseract),
                    "ai_enabled": bool(openai_client),
                    "ai_model": OPENAI_MODEL if openai_client else None,
                    "ai_used": bool(briefing.get("ai_used")),
                    "ai_error": briefing.get("ai_error"),
                },
                "briefing": briefing,
                "sections": dict(sections),
            }

            return JSONResponse(report)

    except Exception as e:
        print("ANALYSE ERROR:")
        traceback.print_exc()
        return JSONResponse({"error": f"Analyse failed: {e}"}, status_code=500)